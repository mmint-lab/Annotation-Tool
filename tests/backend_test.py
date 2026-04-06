import requests
import sys
import json
import io
import csv
from datetime import datetime

class SDOHAPITester:
    def __init__(self, base_url="http://localhost:8000/api"):
        self.base_url = base_url
        self.token = None
        self.admin_token = None
        self.user_id = None
        self.admin_user_id = None
        self.tests_run = 0
        self.tests_passed = 0
        self.test_document_id = None
        self.test_sentence_id = None
        self.created_user_ids = []  # Track created users for cleanup
        self.test_word_resource_id = None  # Track uploaded Word document
        self.test_pdf_resource_id = None   # Track uploaded PDF document

    def run_test(self, name, method, endpoint, expected_status, data=None, files=None):
        """Run a single API test"""
        url = f"{self.base_url}/{endpoint}"
        headers = {}
        
        if self.token:
            headers['Authorization'] = f'Bearer {self.token}'
        
        if not files:
            headers['Content-Type'] = 'application/json'

        self.tests_run += 1
        print(f"\n🔍 Testing {name}...")
        print(f"   URL: {url}")
        
        try:
            if method == 'GET':
                response = requests.get(url, headers=headers)
            elif method == 'POST':
                if files:
                    response = requests.post(url, files=files, headers={k:v for k,v in headers.items() if k != 'Content-Type'})
                else:
                    response = requests.post(url, json=data, headers=headers)
            elif method == 'PUT':
                response = requests.put(url, json=data, headers=headers)
            elif method == 'DELETE':
                response = requests.delete(url, headers=headers)

            success = response.status_code == expected_status
            if success:
                self.tests_passed += 1
                print(f"✅ Passed - Status: {response.status_code}")
                try:
                    response_data = response.json()
                    print(f"   Response: {json.dumps(response_data, indent=2)[:200]}...")
                    return True, response_data
                except:
                    return True, {}
            else:
                print(f"❌ Failed - Expected {expected_status}, got {response.status_code}")
                try:
                    error_data = response.json()
                    print(f"   Error: {error_data}")
                except:
                    print(f"   Error: {response.text}")
                return False, {}

        except Exception as e:
            print(f"❌ Failed - Error: {str(e)}")
            return False, {}

    def test_root_endpoint(self):
        """Test root API endpoint"""
        return self.run_test("Root Endpoint", "GET", "", 200)

    def test_domains_endpoint(self):
        """Test domains endpoint"""
        return self.run_test("Get SDOH Domains", "GET", "domains", 200)

    def test_admin_login(self):
        """Test admin login with predefined credentials"""
        admin_credentials = {
            "email": "admin@sdoh.com",
            "password": "admin123"
        }
        
        success, response = self.run_test(
            "Admin Login",
            "POST",
            "auth/login",
            200,
            data=admin_credentials
        )
        
        if success and 'access_token' in response:
            self.admin_token = response['access_token']
            print(f"   Admin token obtained successfully")
            return True
        return False

    def test_admin_get_current_user(self):
        """Test getting admin user info"""
        # Temporarily switch to admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test("Get Admin User Info", "GET", "auth/me", 200)
        
        if success and response.get('role') == 'admin':
            self.admin_user_id = response.get('id')
            print(f"   Admin user confirmed with role: {response.get('role')}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_get_all_users(self):
        """Test admin endpoint to get all users"""
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test("Admin - Get All Users", "GET", "admin/users", 200)
        
        if success:
            print(f"   Found {len(response)} users in system")
            for user in response[:3]:  # Show first 3 users
                print(f"   - {user.get('full_name')} ({user.get('email')}) - {user.get('role')}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_create_annotator(self):
        """Test admin creating a new annotator account"""
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        timestamp = datetime.now().strftime('%H%M%S')
        new_annotator = {
            "email": f"annotator_{timestamp}@sdoh.com",
            "password": "AnnotatorPass123!",
            "full_name": f"Test Annotator {timestamp}",
            "role": "annotator"
        }
        
        success, response = self.run_test(
            "Admin - Create Annotator",
            "POST",
            "admin/users",
            200,
            data=new_annotator
        )
        
        if success and 'id' in response:
            self.created_user_ids.append(response['id'])
            print(f"   Created annotator with ID: {response['id']}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_create_admin(self):
        """Test admin creating a new admin account"""
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        timestamp = datetime.now().strftime('%H%M%S')
        new_admin = {
            "email": f"admin_{timestamp}@sdoh.com",
            "password": "AdminPass123!",
            "full_name": f"Test Admin {timestamp}",
            "role": "admin"
        }
        
        success, response = self.run_test(
            "Admin - Create Admin",
            "POST",
            "admin/users",
            200,
            data=new_admin
        )
        
        if success and 'id' in response:
            self.created_user_ids.append(response['id'])
            print(f"   Created admin with ID: {response['id']}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_update_user(self):
        """Test admin updating user account"""
        if not self.created_user_ids:
            print("❌ No created users to update")
            return False
            
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        user_id = self.created_user_ids[0]
        update_data = {
            "full_name": "Updated Test User",
            "is_active": False  # Deactivate user
        }
        
        success, response = self.run_test(
            "Admin - Update User",
            "PUT",
            f"admin/users/{user_id}",
            200,
            data=update_data
        )
        
        if success:
            print(f"   Updated user - Active: {response.get('is_active')}")
            print(f"   Updated name: {response.get('full_name')}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_reactivate_user(self):
        """Test admin reactivating user account"""
        if not self.created_user_ids:
            print("❌ No created users to reactivate")
            return False
            
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        user_id = self.created_user_ids[0]
        update_data = {
            "is_active": True  # Reactivate user
        }
        
        success, response = self.run_test(
            "Admin - Reactivate User",
            "PUT",
            f"admin/users/{user_id}",
            200,
            data=update_data
        )
        
        if success:
            print(f"   Reactivated user - Active: {response.get('is_active')}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_document_upload(self):
        """Test admin document upload with project metadata"""
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        # Create test CSV content
        csv_content = self.create_test_csv().getvalue()
        
        # Create file-like object with project metadata
        files = {
            'file': ('admin_test_discharge_summaries.csv', csv_content, 'text/csv'),
            'project_name': (None, 'Admin Test Project'),
            'description': (None, 'Test project uploaded by admin for annotation team')
        }
        
        success, response = self.run_test(
            "Admin - Document Upload with Metadata",
            "POST",
            "documents/upload",
            200,
            files=files
        )
        
        if success and 'id' in response:
            self.test_document_id = response['id']
            print(f"   Document ID: {self.test_document_id}")
            print(f"   Project: {response.get('project_name')}")
            print(f"   Description: {response.get('description')}")
            print(f"   Total sentences: {response.get('total_sentences', 0)}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_delete_document(self):
        """Test admin deleting a document"""
        if not self.test_document_id:
            print("❌ No test document to delete")
            return False
            
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test(
            "Admin - Delete Document",
            "DELETE",
            f"admin/documents/{self.test_document_id}",
            200
        )
        
        if success:
            print(f"   Document deleted successfully")
            self.test_document_id = None  # Clear the ID since document is deleted
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_user_analytics(self):
        """Test admin-only user analytics endpoint"""
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test("Admin - User Analytics", "GET", "admin/analytics/users", 200)
        
        if success:
            print(f"   User analytics retrieved for {len(response)} users")
            # Show sample analytics
            for user_id, analytics in list(response.items())[:2]:
                user_info = analytics.get('user', {})
                print(f"   - {user_info.get('full_name')}: {analytics.get('total_annotations', 0)} annotations")
        
        # Restore original token
        self.token = original_token
        return success

    def test_user_registration(self):
        """Test user registration"""
        timestamp = datetime.now().strftime('%H%M%S')
        test_user_data = {
            "email": f"test_user_{timestamp}@example.com",
            "password": "TestPass123!",
            "full_name": f"Test User {timestamp}"
        }
        
        success, response = self.run_test(
            "User Registration",
            "POST",
            "auth/register",
            200,
            data=test_user_data
        )
        
        if success and 'id' in response:
            self.user_id = response['id']
            # Auto-login after registration
            return self.test_user_login(test_user_data['email'], test_user_data['password'])
        return success

    def test_user_login(self, email, password):
        """Test user login"""
        login_data = {"email": email, "password": password}
        success, response = self.run_test(
            "User Login",
            "POST",
            "auth/login",
            200,
            data=login_data
        )
        
        if success and 'access_token' in response:
            self.token = response['access_token']
            return True
        return False

    def test_get_current_user(self):
        """Test getting current user info"""
        return self.run_test("Get Current User", "GET", "auth/me", 200)

    def test_non_admin_access_denied(self):
        """Test that non-admin users cannot access admin endpoints"""
        if not self.token:
            print("❌ No regular user token available")
            return False
            
        # Try to access admin endpoint with regular user token
        success, response = self.run_test(
            "Non-Admin Access Denied Test",
            "GET",
            "admin/users",
            403  # Expect forbidden
        )
        
        if success:
            print("   ✅ Access properly denied to non-admin user")
        
        return success

    def test_non_admin_upload_denied(self):
        """Test that non-admin users cannot upload documents"""
        if not self.token:
            print("❌ No regular user token available")
            return False
            
        # Create test CSV content
        csv_content = self.create_test_csv().getvalue()
        
        # Create file-like object
        files = {
            'file': ('unauthorized_upload.csv', csv_content, 'text/csv')
        }
        
        success, response = self.run_test(
            "Non-Admin Upload Denied Test",
            "POST",
            "documents/upload",
            403,  # Expect forbidden
            files=files
        )
        
        if success:
            print("   ✅ Upload properly denied to non-admin user")
        
        return success

    def create_test_csv(self):
        """Create a test CSV file with sample medical discharge summaries"""
        csv_content = """patient_id,discharge_summary,notes
1,"Patient is a 45-year-old male with diabetes who was admitted for diabetic ketoacidosis. He reports difficulty affording his insulin medication due to job loss. Lives in a one-bedroom apartment with three family members. Discharge plan includes social work consultation for medication assistance programs.","Financial hardship affecting medication compliance"
2,"67-year-old female with hypertension and COPD. Patient lives alone in a rural area with limited access to healthcare facilities. Nearest pharmacy is 30 miles away. Has high school education and struggles to understand medication instructions. Needs transportation assistance for follow-up appointments.","Geographic and educational barriers to care"
3,"32-year-old mother of two with postpartum depression. Recently moved to new neighborhood and reports feeling isolated. No family support nearby. Concerned about childcare costs affecting ability to attend therapy sessions. Works part-time without health insurance benefits.","Social isolation and economic barriers to mental health care"
"""
        return io.StringIO(csv_content)

    def test_get_documents(self):
        """Test getting all documents"""
        return self.run_test("Get Documents", "GET", "documents", 200)

    def test_get_document_sentences(self):
        """Test getting sentences from a document"""
        if not self.test_document_id:
            print("❌ No test document ID available")
            return False
            
        success, response = self.run_test(
            "Get Document Sentences",
            "GET",
            f"documents/{self.test_document_id}/sentences",
            200
        )
        
        if success and response and len(response) > 0:
            self.test_sentence_id = response[0]['id']
            print(f"   First sentence ID: {self.test_sentence_id}")
            print(f"   First sentence: {response[0]['text'][:100]}...")
        
        return success

    def test_get_tag_structure(self):
        """Test getting the structured tag definitions"""
        return self.run_test("Get Tag Structure", "GET", "tag-structure", 200)

    def test_create_structured_annotation(self):
        """Test creating a structured annotation with new format"""
        if not self.test_sentence_id:
            print("❌ No test sentence ID available")
            return False
            
        # Test structured annotation with multiple tags and valences
        annotation_data = {
            "sentence_id": self.test_sentence_id,
            "tags": [
                {
                    "domain": "Economic Stability",
                    "category": "Employment", 
                    "tag": "Unemployed",
                    "valence": "negative"
                },
                {
                    "domain": "Economic Stability",
                    "category": "Poverty",
                    "tag": "Below Poverty Threshold", 
                    "valence": "negative"
                }
            ],
            "notes": "Patient reports job loss affecting medication access",
            "skipped": False
        }
        
        return self.run_test(
            "Create Structured Annotation",
            "POST",
            "annotations",
            200,
            data=annotation_data
        )

    def test_analytics_overview(self):
        """Test analytics overview endpoint"""
        return self.run_test("Analytics Overview", "GET", "analytics/overview", 200)

    def test_analytics_tag_prevalence(self):
        """Test tag prevalence analytics (updated endpoint)"""
        return self.run_test("Tag Prevalence Analytics", "GET", "analytics/tag-prevalence", 200)

    def test_projects_analytics_endpoint(self):
        """Test /api/analytics/projects endpoint - smoke check"""
        success, response = self.run_test("Projects Analytics JSON", "GET", "analytics/projects", 200)
        
        if success and isinstance(response, list):
            print(f"   Found {len(response)} projects")
            if response:
                project = response[0]
                required_fields = ['project_name', 'documents_count', 'total_sentences', 
                                 'annotated_sentences', 'progress', 'annotators_count', 'last_activity']
                missing_fields = [field for field in required_fields if field not in project]
                if missing_fields:
                    print(f"   ❌ Missing fields: {missing_fields}")
                    return False
                else:
                    print(f"   ✅ All required fields present: {required_fields}")
                    print(f"   Sample project: {project['project_name']} - {project['progress']:.1%} complete")
        
        return success

    def test_projects_chart_authentication(self):
        """Test /api/analytics/projects-chart endpoint authentication scenarios"""
        print("\n🔍 Testing Projects Chart Authentication Scenarios...")
        
        # Store original token
        original_token = self.token
        
        # Test 1: Unauthenticated request should return 401
        self.token = None
        success_unauth, _ = self.run_test(
            "Projects Chart - Unauthenticated (should be 401)",
            "GET",
            "analytics/projects-chart",
            401
        )
        
        # Test 2: Authenticated via Bearer token should return 200 with image/png
        self.token = self.admin_token if self.admin_token else original_token
        success_bearer, response_bearer = self.test_projects_chart_bearer()
        
        # Test 3: Authenticated via token query param should return 200 with image/png
        success_query, response_query = self.test_projects_chart_query_param()
        
        # Test 4: Data integrity test - remaining should be non-negative
        success_integrity = self.test_projects_chart_data_integrity()
        
        # Restore original token
        self.token = original_token
        
        all_passed = success_unauth and success_bearer and success_query and success_integrity
        if all_passed:
            print("   ✅ All authentication scenarios passed")
        else:
            print("   ❌ Some authentication scenarios failed")
        
        return all_passed

    def test_projects_chart_bearer(self):
        """Test projects chart with Bearer token authentication"""
        import requests
        
        url = f"{self.base_url}/analytics/projects-chart"
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        try:
            response = requests.get(url, headers=headers)
            success = response.status_code == 200
            
            if success:
                content_type = response.headers.get('content-type', '')
                content_length = len(response.content)
                
                if 'image/png' in content_type:
                    print(f"   ✅ Bearer Auth - Status: 200, Content-Type: {content_type}, Size: {content_length} bytes")
                    return True, response.content
                else:
                    print(f"   ❌ Bearer Auth - Wrong content type: {content_type}")
                    return False, None
            else:
                print(f"   ❌ Bearer Auth - Status: {response.status_code}")
                return False, None
                
        except Exception as e:
            print(f"   ❌ Bearer Auth - Error: {str(e)}")
            return False, None

    def test_projects_chart_query_param(self):
        """Test projects chart with token query parameter"""
        import requests
        
        url = f"{self.base_url}/analytics/projects-chart?token={self.admin_token}"
        
        try:
            response = requests.get(url)  # No Authorization header
            success = response.status_code == 200
            
            if success:
                content_type = response.headers.get('content-type', '')
                content_length = len(response.content)
                
                if 'image/png' in content_type:
                    print(f"   ✅ Query Param Auth - Status: 200, Content-Type: {content_type}, Size: {content_length} bytes")
                    return True, response.content
                else:
                    print(f"   ❌ Query Param Auth - Wrong content type: {content_type}")
                    return False, None
            else:
                print(f"   ❌ Query Param Auth - Status: {response.status_code}")
                return False, None
                
        except Exception as e:
            print(f"   ❌ Query Param Auth - Error: {str(e)}")
            return False, None

    def test_projects_chart_data_integrity(self):
        """Test data integrity: remaining should be non-negative (total >= annotated)"""
        # First get the projects data to check integrity
        success, projects_data = self.run_test("Projects Data for Integrity Check", "GET", "analytics/projects", 200)
        
        if not success or not projects_data:
            print("   ❌ Could not get projects data for integrity check")
            return False
        
        integrity_issues = []
        for project in projects_data:
            total = project.get('total_sentences', 0)
            annotated = project.get('annotated_sentences', 0)
            remaining = total - annotated
            
            if remaining < 0:
                integrity_issues.append({
                    'project': project.get('project_name', 'Unknown'),
                    'total': total,
                    'annotated': annotated,
                    'remaining': remaining
                })
        
        if integrity_issues:
            print(f"   ❌ Data integrity issues found in {len(integrity_issues)} projects:")
            for issue in integrity_issues:
                print(f"      - {issue['project']}: total={issue['total']}, annotated={issue['annotated']}, remaining={issue['remaining']}")
            return False
        else:
            print(f"   ✅ Data integrity verified: all {len(projects_data)} projects have non-negative remaining sentences")
            # Show sample data
            if projects_data:
                sample = projects_data[0]
                total = sample.get('total_sentences', 0)
                annotated = sample.get('annotated_sentences', 0)
                remaining = total - annotated
                print(f"      Sample: {sample.get('project_name')} - total:{total}, annotated:{annotated}, remaining:{remaining}")
            return True

    def test_admin_delete_user_validation(self):
        """Test user deletion validation (cannot delete self)"""
        if not self.admin_user_id:
            print("❌ No admin user ID available")
            return False
            
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        # Try to delete self - should fail with 400
        success, response = self.run_test(
            "Admin - Delete Self (Should Fail)",
            "DELETE",
            f"admin/users/{self.admin_user_id}",
            400  # Expect bad request
        )
        
        if success:
            print("   ✅ Self-deletion properly prevented")
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_delete_nonexistent_user(self):
        """Test deleting non-existent user"""
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        fake_user_id = "nonexistent-user-id-12345"
        success, response = self.run_test(
            "Admin - Delete Non-existent User",
            "DELETE",
            f"admin/users/{fake_user_id}",
            404  # Expect not found
        )
        
        if success:
            print("   ✅ Non-existent user deletion properly handled")
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_delete_user_with_annotations(self):
        """Test deleting user who has annotations"""
        if not self.created_user_ids:
            print("❌ No created users available for deletion test")
            return False
            
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        user_id = self.created_user_ids[0]
        
        # Get user info before deletion
        success, user_info = self.run_test(
            "Get User Info Before Deletion",
            "GET",
            "admin/users",
            200
        )
        
        if success:
            user_to_delete = next((u for u in user_info if u['id'] == user_id), None)
            if user_to_delete:
                print(f"   User to delete: {user_to_delete['full_name']} ({user_to_delete['email']})")
        
        # Delete the user
        success, response = self.run_test(
            "Admin - Delete User with Annotations",
            "DELETE",
            f"admin/users/{user_id}",
            200
        )
        
        if success:
            print(f"   ✅ User deleted successfully")
            print(f"   Message: {response.get('message', 'N/A')}")
            print(f"   User name: {response.get('user_name', 'N/A')}")
            print(f"   Annotations deleted: {response.get('annotations_deleted', 0)}")
            
            # Remove from created_user_ids since it's deleted
            if user_id in self.created_user_ids:
                self.created_user_ids.remove(user_id)
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_verify_user_deleted(self):
        """Verify that deleted user no longer appears in user list"""
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test(
            "Admin - Verify User List After Deletion",
            "GET",
            "admin/users",
            200
        )
        
        if success:
            print(f"   Current user count: {len(response)}")
            # Check that none of the deleted users appear in the list
            remaining_created_users = [u for u in response if u['id'] in self.created_user_ids]
            if len(remaining_created_users) == 0:
                print("   ✅ Deleted users no longer appear in user list")
            else:
                print(f"   ❌ {len(remaining_created_users)} deleted users still appear in list")
                return False
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_delete_created_users(self):
        """Clean up by deleting remaining created test users"""
        if not self.created_user_ids:
            print("ℹ️  No test users to clean up")
            return True
            
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success_count = 0
        for user_id in self.created_user_ids[:]:  # Create a copy to iterate over
            success, response = self.run_test(
                f"Admin - Delete Test User {user_id[:8]}",
                "DELETE",
                f"admin/users/{user_id}",
                200
            )
            if success:
                success_count += 1
                self.created_user_ids.remove(user_id)
        
        print(f"   Cleaned up {success_count} test users")
        
        # Restore original token
        self.token = original_token
        return True

    def create_test_docx_file(self):
        """Create a test .docx file for preview testing"""
        try:
            from docx import Document
            
            # Create a new document
            doc = Document()
            doc.add_heading('Test Document for Preview', 0)
            
            p = doc.add_paragraph('This is a test document created for testing the Word document preview functionality.')
            p.add_run(' This text should appear in the HTML preview.').bold = True
            
            doc.add_heading('Section 1', level=1)
            doc.add_paragraph('This is the first section with some content.')
            
            doc.add_heading('Section 2', level=2)
            doc.add_paragraph('This is the second section with more content.')
            
            # Add a table
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = 'Header 1'
            table.cell(0, 1).text = 'Header 2'
            table.cell(1, 0).text = 'Data 1'
            table.cell(1, 1).text = 'Data 2'
            
            # Save to BytesIO
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except ImportError:
            # If python-docx is not available, create a minimal docx-like file
            # This is a very basic docx structure
            docx_content = b'PK\x03\x04\x14\x00\x00\x00\x08\x00\x00\x00!\x00'
            return docx_content

    def create_test_pdf_file(self):
        """Create a simple test PDF file"""
        # Minimal PDF content for testing
        pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj

2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj

3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj

4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
72 720 Td
(Test PDF) Tj
ET
endstream
endobj

xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000204 00000 n 
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
297
%%EOF"""
        return pdf_content

    def test_admin_upload_word_document(self):
        """Test admin uploading a Word document for preview testing"""
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        # Create test .docx file
        docx_content = self.create_test_docx_file()
        
        files = {
            'file': ('test_document.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        success, response = self.run_test(
            "Admin - Upload Word Document",
            "POST",
            "admin/resources/upload",
            200,
            files=files
        )
        
        if success and 'id' in response:
            self.test_word_resource_id = response['id']
            print(f"   Word document uploaded with ID: {self.test_word_resource_id}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_admin_upload_pdf_document(self):
        """Test admin uploading a PDF document for negative testing"""
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        # Create test PDF file
        pdf_content = self.create_test_pdf_file()
        
        files = {
            'file': ('test_document.pdf', pdf_content, 'application/pdf')
        }
        
        success, response = self.run_test(
            "Admin - Upload PDF Document",
            "POST",
            "admin/resources/upload",
            200,
            files=files
        )
        
        if success and 'id' in response:
            self.test_pdf_resource_id = response['id']
            print(f"   PDF document uploaded with ID: {self.test_pdf_resource_id}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_word_document_preview_authenticated(self):
        """Test Word document preview with authentication"""
        if not hasattr(self, 'test_word_resource_id') or not self.test_word_resource_id:
            print("❌ No Word document resource ID available")
            return False
        
        # Use admin token for authentication
        original_token = self.token
        self.token = self.admin_token
        
        import requests
        url = f"{self.base_url}/resources/{self.test_word_resource_id}/preview"
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        try:
            response = requests.get(url, headers=headers)
            success = response.status_code == 200
            
            if success:
                content_type = response.headers.get('content-type', '')
                content = response.text
                
                # Verify it's HTML content
                if 'text/html' in content_type:
                    print(f"   ✅ Word Preview - Status: 200, Content-Type: {content_type}")
                    
                    # Check for HTML structure
                    html_checks = [
                        ('<!DOCTYPE html>' in content, 'DOCTYPE declaration'),
                        ('<html>' in content, 'HTML tag'),
                        ('<head>' in content, 'HEAD section'),
                        ('<body>' in content, 'BODY section'),
                        ('font-family:' in content, 'CSS styling'),
                        ('Test Document' in content or 'test' in content.lower(), 'Document content')
                    ]
                    
                    passed_checks = 0
                    for check, description in html_checks:
                        if check:
                            passed_checks += 1
                            print(f"      ✅ {description}")
                        else:
                            print(f"      ❌ {description}")
                    
                    print(f"   HTML validation: {passed_checks}/{len(html_checks)} checks passed")
                    print(f"   Content preview: {content[:200]}...")
                    
                    # Restore original token
                    self.token = original_token
                    return passed_checks >= 4  # At least basic HTML structure
                else:
                    print(f"   ❌ Wrong content type: {content_type}")
                    self.token = original_token
                    return False
            else:
                print(f"   ❌ Word Preview - Status: {response.status_code}")
                try:
                    error_data = response.json()
                    print(f"   Error: {error_data}")
                except:
                    print(f"   Error: {response.text}")
                self.token = original_token
                return False
                
        except Exception as e:
            print(f"   ❌ Word Preview - Error: {str(e)}")
            self.token = original_token
            return False

    def test_word_document_preview_unauthenticated(self):
        """Test Word document preview without authentication (should fail)"""
        if not hasattr(self, 'test_word_resource_id') or not self.test_word_resource_id:
            print("❌ No Word document resource ID available")
            return False
        
        # Remove token for unauthenticated request
        original_token = self.token
        self.token = None
        
        success, response = self.run_test(
            "Word Preview - Unauthenticated (should fail)",
            "GET",
            f"resources/{self.test_word_resource_id}/preview",
            403  # Expect forbidden (not 401)
        )
        
        if success:
            print("   ✅ Authentication properly required for preview")
        
        # Restore original token
        self.token = original_token
        return success

    def test_pdf_document_preview_rejection(self):
        """Test that PDF documents are rejected for preview"""
        if not hasattr(self, 'test_pdf_resource_id') or not self.test_pdf_resource_id:
            print("❌ No PDF document resource ID available")
            return False
        
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test(
            "PDF Preview - Should be rejected",
            "GET",
            f"resources/{self.test_pdf_resource_id}/preview",
            400  # Expect bad request
        )
        
        if success:
            print("   ✅ PDF documents properly rejected for preview")
            if 'detail' in response:
                print(f"   Error message: {response['detail']}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_nonexistent_resource_preview(self):
        """Test preview of non-existent resource"""
        fake_resource_id = "nonexistent-resource-id-12345"
        
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test(
            "Preview Non-existent Resource",
            "GET",
            f"resources/{fake_resource_id}/preview",
            400  # Expect bad request for invalid ID format
        )
        
        if success:
            print("   ✅ Non-existent resource properly handled")
        
        # Restore original token
        self.token = original_token
        return success

    def test_invalid_resource_id_preview(self):
        """Test preview with invalid resource ID format"""
        invalid_resource_id = "invalid-id-format"
        
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test(
            "Preview Invalid Resource ID",
            "GET",
            f"resources/{invalid_resource_id}/preview",
            400  # Expect bad request
        )
        
        if success:
            print("   ✅ Invalid resource ID properly handled")
            if 'detail' in response:
                print(f"   Error message: {response['detail']}")
        
        # Restore original token
        self.token = original_token
        return success

    def test_per_user_csv_export_endpoint(self):
        """Test per-user CSV export endpoint with new columns and features"""
        if not self.test_document_id:
            print("❌ No test document ID available for per-user CSV export test")
            return False
        
        print("\n🔍 Testing Per-User CSV Export Endpoint...")
        
        # Use admin token for testing
        original_token = self.token
        self.token = self.admin_token
        
        # Step 1: Create a tagged annotation with per-tag confidence
        tagged_annotation = {
            "sentence_id": self.test_sentence_id,
            "tags": [
                {
                    "domain": "Economic Stability",
                    "category": "Employment",
                    "tag": "Unemployed",
                    "valence": "negative",
                    "confidence": 4  # Per-tag confidence (1-5 scale)
                },
                {
                    "domain": "Social and Community Context",
                    "category": "Social Cohesion",
                    "tag": "Social Isolation",
                    "valence": "negative", 
                    "confidence": 3  # Different confidence for this tag
                }
            ],
            "notes": "Patient reports job loss and social isolation",
            "skipped": False,
            "duration_ms": 45000  # 45 seconds
        }
        
        success, tagged_response = self.run_test(
            "Create Tagged Annotation with Per-Tag Confidence",
            "POST",
            "annotations",
            200,
            data=tagged_annotation
        )
        
        if not success:
            self.token = original_token
            return False
        
        # Step 2: Create a skipped annotation
        skipped_annotation = {
            "sentence_id": self.test_sentence_id,
            "tags": [],
            "notes": "Sentence not relevant for SDOH annotation",
            "skipped": True,
            "duration_ms": 5000  # 5 seconds
        }
        
        # Get another sentence for skipped annotation
        sentences_success, sentences_response = self.run_test(
            "Get Document Sentences for Skip Test",
            "GET",
            f"documents/{self.test_document_id}/sentences",
            200
        )
        
        if sentences_success and len(sentences_response) > 1:
            # Use second sentence for skipped annotation
            second_sentence_id = sentences_response[1]['id']
            skipped_annotation["sentence_id"] = second_sentence_id
            
            success, skipped_response = self.run_test(
                "Create Skipped Annotation",
                "POST",
                "annotations",
                200,
                data=skipped_annotation
            )
            
            if not success:
                self.token = original_token
                return False
        
        # Step 3: Download per-user CSV and verify content
        import requests
        url = f"{self.base_url}/download/my-annotations-csv/{self.test_document_id}"
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        try:
            response = requests.get(url, headers=headers)
            success = response.status_code == 200
            
            if success:
                content_type = response.headers.get('content-type', '')
                content_disposition = response.headers.get('content-disposition', '')
                
                if 'text/csv' in content_type:
                    print(f"   ✅ CSV Export - Status: 200, Content-Type: {content_type}")
                    print(f"   Content-Disposition: {content_disposition}")
                    
                    # Parse CSV content
                    csv_content = response.text
                    csv_lines = csv_content.strip().split('\n')
                    
                    if len(csv_lines) > 0:
                        # Check header row
                        header = csv_lines[0]
                        expected_columns = [
                            "document_id", "sentence_id", "subject_id", "row_index", 
                            "sentence_index", "sentence_text", "tag_domain", "tag_category", 
                            "tag", "valence", "confidence", "notes", "is_skipped", 
                            "timestamp", "duration_ms"
                        ]
                        
                        print(f"   CSV Header: {header}")
                        
                        # Verify all expected columns are present
                        header_columns = [col.strip('"') for col in header.split(',')]
                        missing_columns = [col for col in expected_columns if col not in header_columns]
                        
                        if missing_columns:
                            print(f"   ❌ Missing columns: {missing_columns}")
                            self.token = original_token
                            return False
                        else:
                            print(f"   ✅ All required columns present: {len(expected_columns)} columns")
                        
                        # Parse CSV data
                        import csv as csv_module
                        from io import StringIO
                        
                        csv_reader = csv_module.DictReader(StringIO(csv_content))
                        rows = list(csv_reader)
                        
                        print(f"   CSV contains {len(rows)} data rows")
                        
                        # Verify tagged annotations have per-tag confidence
                        tagged_rows = [row for row in rows if row['is_skipped'].lower() == 'false' and row['tag_domain']]
                        skipped_rows = [row for row in rows if row['is_skipped'].lower() == 'true']
                        
                        print(f"   Found {len(tagged_rows)} tagged annotation rows")
                        print(f"   Found {len(skipped_rows)} skipped annotation rows")
                        
                        # Verify tagged annotations
                        confidence_check = True
                        for row in tagged_rows:
                            if not row['confidence'] or row['confidence'] == '':
                                print(f"   ❌ Tagged row missing confidence: {row['tag']}")
                                confidence_check = False
                            else:
                                print(f"   ✅ Tagged row has confidence {row['confidence']}: {row['tag_domain']}:{row['tag']}")
                        
                        # Verify skipped annotations
                        skipped_check = True
                        for row in skipped_rows:
                            if row['is_skipped'].lower() != 'true':
                                print(f"   ❌ Skipped row not marked as skipped")
                                skipped_check = False
                            else:
                                print(f"   ✅ Skipped annotation properly marked: is_skipped=TRUE")
                        
                        # Verify timestamp column
                        timestamp_check = True
                        for row in rows:
                            if not row['timestamp'] or row['timestamp'] == '':
                                print(f"   ❌ Row missing timestamp")
                                timestamp_check = False
                                break
                        
                        if timestamp_check:
                            print(f"   ✅ All rows have timestamp values")
                        
                        # Verify duration_ms column
                        duration_check = True
                        for row in rows:
                            if row['duration_ms'] == '':
                                print(f"   ⚠️  Row missing duration_ms (acceptable)")
                            else:
                                try:
                                    duration = int(row['duration_ms'])
                                    if duration > 0:
                                        print(f"   ✅ Row has valid duration_ms: {duration}ms")
                                except ValueError:
                                    print(f"   ❌ Invalid duration_ms value: {row['duration_ms']}")
                                    duration_check = False
                        
                        all_checks_passed = confidence_check and skipped_check and timestamp_check and duration_check
                        
                        if all_checks_passed:
                            print("   ✅ Per-User CSV Export: ALL CHECKS PASSED")
                            self.token = original_token
                            return True
                        else:
                            print("   ❌ Per-User CSV Export: Some checks failed")
                            self.token = original_token
                            return False
                    else:
                        print("   ❌ CSV content is empty")
                        self.token = original_token
                        return False
                else:
                    print(f"   ❌ Wrong content type: {content_type}")
                    self.token = original_token
                    return False
            else:
                print(f"   ❌ CSV Export failed - Status: {response.status_code}")
                try:
                    error_data = response.json()
                    print(f"   Error: {error_data}")
                except:
                    print(f"   Error: {response.text}")
                self.token = original_token
                return False
                
        except Exception as e:
            print(f"   ❌ CSV Export - Error: {str(e)}")
            self.token = original_token
            return False

    def test_per_user_paragraph_export_endpoint(self):
        """Test per-user paragraph export endpoint with new format"""
        if not self.test_document_id:
            print("❌ No test document ID available for per-user paragraph export test")
            return False
        
        print("\n🔍 Testing Per-User Paragraph Export Endpoint...")
        
        # Use admin token for testing
        original_token = self.token
        self.token = self.admin_token
        
        # Download per-user paragraph export
        import requests
        url = f"{self.base_url}/download/my-annotated-paragraphs/{self.test_document_id}"
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        try:
            response = requests.get(url, headers=headers)
            success = response.status_code == 200
            
            if success:
                content_type = response.headers.get('content-type', '')
                content_disposition = response.headers.get('content-disposition', '')
                
                if 'text/csv' in content_type:
                    print(f"   ✅ Paragraph Export - Status: 200, Content-Type: {content_type}")
                    print(f"   Content-Disposition: {content_disposition}")
                    
                    # Parse CSV content
                    csv_content = response.text
                    csv_lines = csv_content.strip().split('\n')
                    
                    if len(csv_lines) > 0:
                        # Check header row
                        header = csv_lines[0]
                        expected_columns = ["row_index", "subject_id", "annotated_paragraph_text"]
                        
                        print(f"   CSV Header: {header}")
                        
                        # Verify expected columns
                        header_columns = [col.strip('"') for col in header.split(',')]
                        missing_columns = [col for col in expected_columns if col not in header_columns]
                        
                        if missing_columns:
                            print(f"   ❌ Missing columns: {missing_columns}")
                            self.token = original_token
                            return False
                        else:
                            print(f"   ✅ All required columns present: {expected_columns}")
                        
                        # Parse CSV data
                        import csv as csv_module
                        from io import StringIO
                        
                        csv_reader = csv_module.DictReader(StringIO(csv_content))
                        rows = list(csv_reader)
                        
                        print(f"   CSV contains {len(rows)} paragraph rows")
                        
                        # Check for new tag format with timestamp and confidence
                        format_checks = {
                            'tags_with_confidence': False,
                            'tags_with_timestamp': False,
                            'skipped_with_timestamp': False
                        }
                        
                        for row in rows:
                            paragraph_text = row['annotated_paragraph_text']
                            
                            # Check for tag format: Domain:Category:Tag(valence,conf=X)@User@Timestamp
                            if '[Tags:' in paragraph_text:
                                print(f"   Found tagged paragraph: {paragraph_text[:100]}...")
                                
                                # Check for confidence in tag format
                                if 'conf=' in paragraph_text:
                                    format_checks['tags_with_confidence'] = True
                                    print("   ✅ Found per-tag confidence in format")
                                
                                # Check for timestamp in tag format
                                if '@' in paragraph_text and paragraph_text.count('@') >= 2:
                                    format_checks['tags_with_timestamp'] = True
                                    print("   ✅ Found timestamp in tag format")
                            
                            # Check for skipped format: [SKIPPED@User@Timestamp]
                            if '[SKIPPED@' in paragraph_text:
                                format_checks['skipped_with_timestamp'] = True
                                print(f"   ✅ Found skipped annotation with timestamp: {paragraph_text}")
                        
                        # Verify format requirements
                        all_format_checks_passed = True
                        
                        if not format_checks['tags_with_confidence']:
                            print("   ⚠️  No tagged annotations with confidence found (may be acceptable if no tagged annotations exist)")
                        
                        if not format_checks['tags_with_timestamp']:
                            print("   ⚠️  No tagged annotations with timestamp found (may be acceptable if no tagged annotations exist)")
                        
                        if not format_checks['skipped_with_timestamp']:
                            print("   ⚠️  No skipped annotations with timestamp found (may be acceptable if no skipped annotations exist)")
                        
                        # At least one format should be present if we have annotations
                        if any(format_checks.values()):
                            print("   ✅ Per-User Paragraph Export: Format checks passed")
                            self.token = original_token
                            return True
                        else:
                            print("   ⚠️  No annotations found in expected format (may be acceptable for empty document)")
                            self.token = original_token
                            return True  # Consider this a pass if document has no annotations
                    else:
                        print("   ❌ CSV content is empty")
                        self.token = original_token
                        return False
                else:
                    print(f"   ❌ Wrong content type: {content_type}")
                    self.token = original_token
                    return False
            else:
                print(f"   ❌ Paragraph Export failed - Status: {response.status_code}")
                try:
                    error_data = response.json()
                    print(f"   Error: {error_data}")
                except:
                    print(f"   Error: {response.text}")
                self.token = original_token
                return False
                
        except Exception as e:
            print(f"   ❌ Paragraph Export - Error: {str(e)}")
            self.token = original_token
            return False

    def test_documents_no_project_name_badge(self):
        """Test that documents do not have project_name badge in response"""
        print("\n🔍 Testing Documents API - No Project Name Badge...")
        
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test(
            "Get Documents - Check No Project Name Badge",
            "GET",
            "documents",
            200
        )
        
        if success and isinstance(response, list):
            print(f"   Found {len(response)} documents")
            
            project_name_found = False
            for doc in response:
                if 'project_name' in doc and doc['project_name']:
                    project_name_found = True
                    print(f"   ⚠️  Document has project_name: {doc.get('filename')} - {doc['project_name']}")
                else:
                    print(f"   ✅ Document without project_name badge: {doc.get('filename')}")
            
            if not project_name_found:
                print("   ✅ No documents have project_name badge (as expected)")
                self.token = original_token
                return True
            else:
                print("   ❌ Some documents still have project_name badge")
                self.token = original_token
                return False
        else:
            print("   ❌ Failed to get documents or invalid response format")
            self.token = original_token
            return False

    def test_domain_tag_stats_endpoint(self):
        """Test /api/analytics/domain-tag-stats endpoint"""
        print("\n🔍 Testing Domain Tag Stats Endpoint...")
        
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test(
            "Domain Tag Stats Analytics",
            "GET",
            "analytics/domain-tag-stats",
            200
        )
        
        if success and isinstance(response, dict):
            # Check required fields
            required_fields = ['domain_totals', 'tag_details', 'domains']
            missing_fields = [field for field in required_fields if field not in response]
            
            if missing_fields:
                print(f"   ❌ Missing required fields: {missing_fields}")
                self.token = original_token
                return False
            
            print(f"   ✅ All required fields present: {required_fields}")
            
            # Verify domain_totals structure
            domain_totals = response.get('domain_totals', {})
            if isinstance(domain_totals, dict):
                print(f"   ✅ domain_totals is object with {len(domain_totals)} domains")
                for domain, count in list(domain_totals.items())[:3]:
                    print(f"      - {domain}: {count} tags")
            else:
                print(f"   ❌ domain_totals should be object, got {type(domain_totals)}")
                self.token = original_token
                return False
            
            # Verify tag_details structure (nested object)
            tag_details = response.get('tag_details', {})
            if isinstance(tag_details, dict):
                print(f"   ✅ tag_details is nested object with {len(tag_details)} domains")
                for domain, categories in list(tag_details.items())[:2]:
                    if isinstance(categories, dict):
                        print(f"      - {domain}: {len(categories)} categories")
                        for category, tags in list(categories.items())[:2]:
                            if isinstance(tags, dict):
                                print(f"        - {category}: {len(tags)} tags")
                            else:
                                print(f"   ❌ Tags should be object, got {type(tags)}")
                                self.token = original_token
                                return False
                    else:
                        print(f"   ❌ Categories should be object, got {type(categories)}")
                        self.token = original_token
                        return False
            else:
                print(f"   ❌ tag_details should be object, got {type(tag_details)}")
                self.token = original_token
                return False
            
            # Verify domains array
            domains = response.get('domains', [])
            if isinstance(domains, list):
                print(f"   ✅ domains is array with {len(domains)} SDOH domain names")
                expected_domains = ["Economic Stability", "Education Access and Quality", 
                                  "Health Care Access and Quality", "Neighborhood and Built Environment",
                                  "Social and Community Context"]
                for domain in expected_domains:
                    if domain in domains:
                        print(f"      ✅ Found expected domain: {domain}")
                    else:
                        print(f"      ⚠️  Missing expected domain: {domain}")
            else:
                print(f"   ❌ domains should be array, got {type(domains)}")
                self.token = original_token
                return False
            
            print("   ✅ Domain Tag Stats endpoint: ALL CHECKS PASSED")
            self.token = original_token
            return True
        else:
            print(f"   ❌ Expected dict response, got {type(response)}")
            self.token = original_token
            return False

    def test_domain_chart_endpoint(self):
        """Test /api/analytics/domain-chart/{domain_name} endpoint"""
        print("\n🔍 Testing Domain Chart Endpoint...")
        
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        # Test with "Economic Stability" domain
        domain_name = "Economic Stability"
        
        import requests
        url = f"{self.base_url}/analytics/domain-chart/{domain_name}"
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        try:
            response = requests.get(url, headers=headers)
            success = response.status_code == 200
            
            if success:
                content_type = response.headers.get('content-type', '')
                content_length = len(response.content)
                
                if 'image/png' in content_type:
                    print(f"   ✅ Domain Chart - Status: 200, Content-Type: {content_type}")
                    print(f"   ✅ PNG image returned with size: {content_length} bytes")
                    print(f"   ✅ Domain tested: {domain_name}")
                    
                    # Test authentication requirement
                    unauth_response = requests.get(url)  # No auth header
                    if unauth_response.status_code == 401:
                        print(f"   ✅ Authentication properly required (401 without token)")
                    else:
                        print(f"   ⚠️  Expected 401 without auth, got {unauth_response.status_code}")
                    
                    self.token = original_token
                    return True
                else:
                    print(f"   ❌ Expected image/png, got content-type: {content_type}")
                    self.token = original_token
                    return False
            else:
                print(f"   ❌ Domain Chart failed - Status: {response.status_code}")
                try:
                    error_data = response.json()
                    print(f"   Error: {error_data}")
                except:
                    print(f"   Error: {response.text}")
                self.token = original_token
                return False
                
        except Exception as e:
            print(f"   ❌ Domain Chart - Error: {str(e)}")
            self.token = original_token
            return False

    def test_all_documents_user_progress_endpoint(self):
        """Test /api/analytics/all-documents-user-progress endpoint"""
        print("\n🔍 Testing All Documents User Progress Endpoint...")
        
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        success, response = self.run_test(
            "All Documents User Progress Analytics",
            "GET",
            "analytics/all-documents-user-progress",
            200
        )
        
        if success and isinstance(response, list):
            print(f"   ✅ Response is array with {len(response)} documents")
            
            if len(response) > 0:
                # Check structure of first document
                doc = response[0]
                required_fields = ['filename', 'total_sentences', 'user_progress']
                missing_fields = [field for field in required_fields if field not in doc]
                
                if missing_fields:
                    print(f"   ❌ Missing required fields in document: {missing_fields}")
                    self.token = original_token
                    return False
                
                print(f"   ✅ Document has required fields: {required_fields}")
                print(f"   Document: {doc.get('filename')} - {doc.get('total_sentences')} sentences")
                
                # Check user_progress structure
                user_progress = doc.get('user_progress', [])
                if isinstance(user_progress, list):
                    print(f"   ✅ user_progress is array with {len(user_progress)} users")
                    
                    if len(user_progress) > 0:
                        user = user_progress[0]
                        required_user_fields = ['user_name', 'annotated', 'total', 'progress']
                        missing_user_fields = [field for field in required_user_fields if field not in user]
                        
                        if missing_user_fields:
                            print(f"   ❌ Missing required fields in user_progress: {missing_user_fields}")
                            self.token = original_token
                            return False
                        
                        print(f"   ✅ User progress has required fields: {required_user_fields}")
                        print(f"   Sample user: {user.get('user_name')} - {user.get('annotated')}/{user.get('total')} ({user.get('progress'):.1%})")
                    else:
                        print(f"   ⚠️  No user progress data (acceptable if no annotations exist)")
                else:
                    print(f"   ❌ user_progress should be array, got {type(user_progress)}")
                    self.token = original_token
                    return False
            else:
                print(f"   ⚠️  No documents found (acceptable if no documents exist)")
            
            print("   ✅ All Documents User Progress endpoint: ALL CHECKS PASSED")
            self.token = original_token
            return True
        else:
            print(f"   ❌ Expected array response, got {type(response)}")
            self.token = original_token
            return False

    def test_activity_log_with_user_filter_endpoint(self):
        """Test /api/admin/download/activity-log endpoint with user filter"""
        print("\n🔍 Testing Activity Log with User Filter Endpoint...")
        
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        # Test 1: Download all activities (no user filter)
        import requests
        url = f"{self.base_url}/admin/download/activity-log"
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        try:
            response = requests.get(url, headers=headers)
            success = response.status_code == 200
            
            if success:
                content_type = response.headers.get('content-type', '')
                content_disposition = response.headers.get('content-disposition', '')
                
                if 'text/csv' in content_type:
                    print(f"   ✅ Activity Log (All) - Status: 200, Content-Type: {content_type}")
                    print(f"   Content-Disposition: {content_disposition}")
                    
                    # Parse CSV content
                    csv_content = response.text
                    csv_lines = csv_content.strip().split('\n')
                    
                    if len(csv_lines) > 0:
                        header = csv_lines[0]
                        expected_columns = ["timestamp", "user_id", "user_name", "document_id", 
                                          "sentence_id", "action_type", "metadata"]
                        
                        print(f"   CSV Header: {header}")
                        
                        # Verify columns
                        header_columns = [col.strip('"') for col in header.split(',')]
                        missing_columns = [col for col in expected_columns if col not in header_columns]
                        
                        if missing_columns:
                            print(f"   ❌ Missing columns: {missing_columns}")
                        else:
                            print(f"   ✅ All required columns present: {len(expected_columns)} columns")
                        
                        print(f"   Activity log contains {len(csv_lines) - 1} data rows")
                    else:
                        print("   ⚠️  Activity log is empty (acceptable if no activities logged)")
                else:
                    print(f"   ❌ Expected text/csv, got content-type: {content_type}")
                    self.token = original_token
                    return False
            else:
                print(f"   ❌ Activity Log (All) failed - Status: {response.status_code}")
                self.token = original_token
                return False
            
            # Test 2: Download activities for specific user (if we have a user ID)
            if hasattr(self, 'admin_user_id') and self.admin_user_id:
                url_with_filter = f"{self.base_url}/admin/download/activity-log?user_id={self.admin_user_id}"
                
                response_filtered = requests.get(url_with_filter, headers=headers)
                success_filtered = response_filtered.status_code == 200
                
                if success_filtered:
                    content_type_filtered = response_filtered.headers.get('content-type', '')
                    
                    if 'text/csv' in content_type_filtered:
                        print(f"   ✅ Activity Log (Filtered) - Status: 200, Content-Type: {content_type_filtered}")
                        
                        csv_content_filtered = response_filtered.text
                        csv_lines_filtered = csv_content_filtered.strip().split('\n')
                        
                        print(f"   Filtered activity log contains {len(csv_lines_filtered) - 1} data rows")
                        print(f"   ✅ User filter working (user_id={self.admin_user_id})")
                    else:
                        print(f"   ❌ Expected text/csv for filtered log, got: {content_type_filtered}")
                        self.token = original_token
                        return False
                else:
                    print(f"   ❌ Activity Log (Filtered) failed - Status: {response_filtered.status_code}")
                    self.token = original_token
                    return False
            else:
                print("   ⚠️  No admin user ID available for filter test")
            
            print("   ✅ Activity Log with User Filter endpoint: ALL CHECKS PASSED")
            self.token = original_token
            return True
                
        except Exception as e:
            print(f"   ❌ Activity Log - Error: {str(e)}")
            self.token = original_token
            return False

    def test_resource_preview_excel_endpoint(self):
        """Test /api/resources/{id}/preview endpoint for Excel files"""
        print("\n🔍 Testing Resource Preview for Excel Files...")
        
        # Use admin token
        original_token = self.token
        self.token = self.admin_token
        
        # First, upload an Excel file for testing
        excel_content = self.create_test_excel_file()
        
        files = {
            'file': ('test_spreadsheet.xlsx', excel_content, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        }
        
        success, upload_response = self.run_test(
            "Admin - Upload Excel File for Preview Test",
            "POST",
            "admin/resources/upload",
            200,
            files=files
        )
        
        if not success or 'id' not in upload_response:
            print("   ❌ Failed to upload Excel file for testing")
            self.token = original_token
            return False
        
        excel_resource_id = upload_response['id']
        print(f"   Excel file uploaded with ID: {excel_resource_id}")
        
        # Test Excel preview
        import requests
        url = f"{self.base_url}/resources/{excel_resource_id}/preview"
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        try:
            response = requests.get(url, headers=headers)
            success = response.status_code == 200
            
            if success:
                content_type = response.headers.get('content-type', '')
                content = response.text
                
                if 'text/html' in content_type:
                    print(f"   ✅ Excel Preview - Status: 200, Content-Type: {content_type}")
                    
                    # Check for HTML table structure
                    html_checks = [
                        ('<table' in content, 'HTML table tag'),
                        ('<tr' in content, 'Table rows'),
                        ('<td' in content or '<th' in content, 'Table cells'),
                        ('Header 1' in content or 'Data' in content, 'Table content'),
                        ('<!DOCTYPE html>' in content or '<html>' in content, 'HTML structure')
                    ]
                    
                    passed_checks = 0
                    for check, description in html_checks:
                        if check:
                            passed_checks += 1
                            print(f"      ✅ {description}")
                        else:
                            print(f"      ❌ {description}")
                    
                    print(f"   HTML table validation: {passed_checks}/{len(html_checks)} checks passed")
                    
                    # Check for first 10 rows limitation
                    row_count = content.count('<tr')
                    if row_count <= 11:  # Header + max 10 data rows
                        print(f"   ✅ Row limit respected: {row_count} rows (including header)")
                    else:
                        print(f"   ⚠️  More than 10 rows found: {row_count} (may be acceptable)")
                    
                    print(f"   Content preview: {content[:200]}...")
                    
                    self.token = original_token
                    return passed_checks >= 3  # At least basic table structure
                else:
                    print(f"   ❌ Expected text/html, got content-type: {content_type}")
                    self.token = original_token
                    return False
            else:
                print(f"   ❌ Excel Preview failed - Status: {response.status_code}")
                try:
                    error_data = response.json()
                    print(f"   Error: {error_data}")
                except:
                    print(f"   Error: {response.text}")
                self.token = original_token
                return False
                
        except Exception as e:
            print(f"   ❌ Excel Preview - Error: {str(e)}")
            self.token = original_token
            return False

    def create_test_excel_file(self):
        """Create a test Excel file for preview testing"""
        try:
            import openpyxl
            from io import BytesIO
            
            # Create a new workbook and worksheet
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Test Sheet"
            
            # Add headers
            headers = ["Header 1", "Header 2", "Header 3", "Header 4"]
            for col, header in enumerate(headers, 1):
                ws.cell(row=1, column=col, value=header)
            
            # Add test data (15 rows to test the 10-row limit)
            for row in range(2, 17):
                for col in range(1, 5):
                    ws.cell(row=row, column=col, value=f"Data {row}-{col}")
            
            # Save to BytesIO
            buffer = BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            # If openpyxl is not available, create a minimal xlsx-like file
            # This is a very basic xlsx structure for testing
            xlsx_content = b'PK\x03\x04\x14\x08!'
            return xlsx_content

    def run_all_tests(self):
        """Run all API tests including comprehensive admin functionality"""
        print("🚀 Starting SDOH API Tests with Admin Functionality")
        print("=" * 60)
        
        # Basic endpoint tests
        self.test_root_endpoint()
        self.test_domains_endpoint()
        self.test_get_tag_structure()
        
        # === ADMIN FUNCTIONALITY TESTS ===
        print("\n" + "=" * 30 + " ADMIN TESTS " + "=" * 30)
        
        # Admin authentication
        if not self.test_admin_login():
            print("❌ Admin login failed, stopping admin tests")
            return False
            
        self.test_admin_get_current_user()
        
        # Admin user management
        self.test_admin_get_all_users()
        self.test_admin_create_annotator()
        self.test_admin_create_admin()
        self.test_admin_update_user()
        self.test_admin_reactivate_user()
        
        # === USER DELETION TESTS ===
        print("\n" + "=" * 25 + " USER DELETION TESTS " + "=" * 25)
        self.test_admin_delete_user_validation()
        self.test_admin_delete_nonexistent_user()
        self.test_admin_delete_user_with_annotations()
        self.test_admin_verify_user_deleted()
        
        # Admin document management
        self.test_admin_document_upload()
        if self.test_document_id:
            self.test_get_documents()
            self.test_get_document_sentences()
            if self.test_sentence_id:
                self.test_create_structured_annotation()
        
        # Admin analytics
        self.test_admin_user_analytics()
        
        # Admin document deletion
        if self.test_document_id:
            self.test_admin_delete_document()
        
        # === WORD DOCUMENT PREVIEW TESTS ===
        print("\n" + "=" * 20 + " WORD DOCUMENT PREVIEW TESTS " + "=" * 20)
        
        # Upload test documents for preview testing
        self.test_admin_upload_word_document()
        self.test_admin_upload_pdf_document()
        
        # Test Word document preview functionality
        if hasattr(self, 'test_word_resource_id') and self.test_word_resource_id:
            self.test_word_document_preview_authenticated()
            self.test_word_document_preview_unauthenticated()
        
        # Test file type validation
        if hasattr(self, 'test_pdf_resource_id') and self.test_pdf_resource_id:
            self.test_pdf_document_preview_rejection()
        
        # Test error handling
        self.test_nonexistent_resource_preview()
        self.test_invalid_resource_id_preview()
        
        # === REGULAR USER TESTS ===
        print("\n" + "=" * 25 + " REGULAR USER TESTS " + "=" * 25)
        
        # Regular user authentication
        if not self.test_user_registration():
            print("❌ Registration failed, stopping regular user tests")
        else:
            self.test_get_current_user()
            
            # Test access control
            self.test_non_admin_access_denied()
            self.test_non_admin_upload_denied()
        
        # Analytics tests (available to all users)
        self.test_analytics_overview()
        self.test_analytics_tag_prevalence()
        
        # === FOCUSED PROJECTS ANALYTICS TESTS ===
        print("\n" + "=" * 20 + " PROJECTS ANALYTICS TESTS " + "=" * 20)
        self.test_projects_analytics_endpoint()
        self.test_projects_chart_authentication()
        
        # === NEW ADMIN ANALYTICS ENDPOINTS TESTS ===
        print("\n" + "=" * 15 + " NEW ADMIN ANALYTICS ENDPOINTS TESTS " + "=" * 15)
        self.test_domain_tag_stats_endpoint()
        self.test_domain_chart_endpoint()
        self.test_all_documents_user_progress_endpoint()
        self.test_activity_log_with_user_filter_endpoint()
        self.test_resource_preview_excel_endpoint()
        
        # === CLEANUP ===
        print("\n" + "=" * 30 + " CLEANUP " + "=" * 30)
        self.test_admin_delete_created_users()
        
        # Print final results
        print("\n" + "=" * 60)
        print(f"📊 Test Results: {self.tests_passed}/{self.tests_run} tests passed")
        
        if self.tests_passed == self.tests_run:
            print("🎉 All tests passed!")
            return True
        else:
            print(f"⚠️  {self.tests_run - self.tests_passed} tests failed")
            return False

def main():
    tester = SDOHAPITester()
    success = tester.run_all_tests()
    return 0 if success else 1

if __name__ == "__main__":
    sys.exit(main())